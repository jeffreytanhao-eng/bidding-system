
import re
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class TenderParser:
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except ImportError:
            logger.warning("pdfplumber not installed, skipping PDF parsing")
            return ""
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return ""

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            return text
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX parsing")
            return ""
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return ""

    @staticmethod
    def extract_text_from_excel(file_content: bytes) -> str:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(file_content))
            text = ""
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell else "" for cell in row])
                    text += row_text + "\n"
            return text
        except ImportError:
            logger.warning("openpyxl not installed, skipping Excel parsing")
            return ""
        except Exception as e:
            logger.error(f"Error parsing Excel: {e}")
            return ""

    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> str:
        ext = filename.lower().split(".")[-1] if "." in filename else ""
        if ext == "pdf":
            return TenderParser.extract_text_from_pdf(file_content)
        elif ext in ("docx", "doc"):
            return TenderParser.extract_text_from_docx(file_content)
        elif ext in ("xlsx", "xls"):
            return TenderParser.extract_text_from_excel(file_content)
        elif ext == "txt":
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                return file_content.decode("gbk", errors="ignore")
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return ""

    @staticmethod
    def parse_tender_info(text: str) -> dict:
        info = {
            "tender_no": None,
            "title": None,
            "budget": None,
            "deadline": None,
            "procurement_method": None,
            "description": None,
            "extracted_text": text[:2000] if text else None,
        }

        if not text:
            return info

        tender_no_patterns = [
            r"(?:招标|项目|采购)[编号：:]\s*([A-Za-z0-9\-]+)",
            r"(?:编号|项目编号|招标编号)[：:]\s*([A-Za-z0-9\-]+)",
            r"(?:ZB|CG|XM|XJ)[\-]\d{4}[\-]\d+",
            r"\d{4}[\-]\d{2,4}[\-]\d+",
        ]
        for pattern in tender_no_patterns:
            match = re.search(pattern, text)
            if match:
                info["tender_no"] = match.group(1) if match.lastindex else match.group(0)
                break

        title_patterns = [
            r"(?:项目名称|招标项目|工程名称|采购项目)[：:]\s*(.+?)(?:\n|$)",
            r"(?:关于|关于对).{2,50}(?:采购|招标|项目)",
        ]
        for pattern in title_patterns:
            match = re.search(pattern, text)
            if match:
                title = match.group(1).strip() if match.lastindex else match.group(0).strip()
                if len(title) > 5:
                    info["title"] = title[:200]
                    break

        budget_patterns = [
            r"(?:预算金额|控制价|最高限价|预算|采购预算|项目预算)[：:（(]?\s*([\d,，.]+)\s*(?:万?元)",
            r"(?:预算金额|控制价|最高限价|预算|采购预算|项目预算)[：:（(]?\s*人民币\s*([\d,，.]+)\s*(?:万?元)",
            r"([\d,，.]+)\s*(?:万?元)(?:.*(?:预算|控制价|限价))",
        ]
        for pattern in budget_patterns:
            match = re.search(pattern, text)
            if match:
                budget_str = match.group(1).replace(",", "").replace("，", "")
                try:
                    budget_val = float(budget_str)
                    if "万" in match.group(0):
                        budget_val *= 10000
                    info["budget"] = budget_val
                except ValueError:
                    pass
                break

        deadline_patterns = [
            r"(?:投标截止|递交截止|截止时间|开标时间)[：:]\s*(\d{4}[\-年/]\d{1,2}[\-月/]\d{1,2}[\s日]*\d{0,2}[:时]?\d{0,2})",
            r"(\d{4}[\-年/]\d{1,2}[\-月/]\d{1,2}[\s日]*\d{0,2}[:时]?\d{0,2}).*?(?:截止|开标)",
        ]
        for pattern in deadline_patterns:
            match = re.search(pattern, text)
            if match:
                deadline_str = match.group(1)
                deadline_str = re.sub(r"[年月]", "-", deadline_str)
                deadline_str = re.sub(r"日", " ", deadline_str)
                deadline_str = re.sub(r"时", ":", deadline_str)
                deadline_str = deadline_str.strip()
                info["deadline"] = deadline_str
                break

        method_patterns = [
            (r"公开招标", "public_bidding"),
            (r"邀请招标", "invited_bidding"),
            (r"竞争性谈判", "competitive_negotiation"),
            (r"单一来源", "single_source"),
            (r"竞争性磋商", "competitive_consultation"),
            (r"询价", "inquiry"),
        ]
        for pattern, method in method_patterns:
            if re.search(pattern, text):
                info["procurement_method"] = method
                break

        lines = text.split("\n")
        desc_lines = []
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if any(kw in stripped for kw in ["项目概况", "采购需求", "项目简介", "一、项目"]):
                for i in range(idx + 1, min(idx + 10, len(lines))):
                    l = lines[i].strip()
                    if l and not l.startswith(("二", "三", "四", "1.", "2.")):
                        desc_lines.append(l)
                    elif l.startswith(("二", "三", "四")):
                        break
                break
        if desc_lines:
            info["description"] = "\n".join(desc_lines)[:1000]

        return info


tender_parser = TenderParser()
